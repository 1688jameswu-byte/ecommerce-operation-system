from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


OUT_DIR = Path("artifacts")
DOCX_PATH = OUT_DIR / "北师版四年级下册_运算定律_A4讲义.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=60, start=90, bottom=60, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D7DEE8", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_fixed_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_cm / 2.54 * 1440)))


def format_run(run, size=8.0, bold=False, color="1F2937"):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def paragraph(cell, text="", size=8.0, bold=False, color="1F2937", before=0, after=1.5, align=None):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.05
    if align:
        p.alignment = align
    r = p.add_run(text)
    format_run(r, size=size, bold=bold, color=color)
    return p


def add_mixed_para(cell, parts, size=8.0, before=0, after=1.5):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.05
    for text, bold, color in parts:
        r = p.add_run(text)
        format_run(r, size=size, bold=bold, color=color)
    return p


def section_label(cell, text):
    p = paragraph(cell, text, size=8.5, bold=True, color="FFFFFF", before=3, after=2)
    set_cell_shading(cell, "FFFFFF")
    return p


def add_rule_card(cell, title, rule, formula, point, fill="F6F9FC"):
    table = cell.add_table(rows=1, cols=1)
    set_table_borders(table, color="DCE3EC", size="3")
    inner = table.cell(0, 0)
    set_cell_shading(inner, fill)
    set_cell_margins(inner, top=55, start=80, bottom=55, end=80)
    paragraph(inner, title, size=8.5, bold=True, color="1F4D78", after=1)
    add_mixed_para(inner, [("规则：", True, "374151"), (rule, False, "1F2937")], size=7.4, after=0.8)
    add_mixed_para(inner, [("公式：", True, "374151"), (formula, True, "B91C1C")], size=7.7, after=0.8)
    add_mixed_para(inner, [("变化要点：", True, "374151"), (point, False, "1F2937")], size=7.4, after=0)


def add_small_box(cell, title, items, fill="FFF7ED", accent="C2410C"):
    table = cell.add_table(rows=1, cols=1)
    set_table_borders(table, color="F6D3B7", size="4")
    inner = table.cell(0, 0)
    set_cell_shading(inner, fill)
    set_cell_margins(inner, top=60, start=95, bottom=60, end=95)
    paragraph(inner, title, size=8.5, bold=True, color=accent, after=1.5)
    for item in items:
        paragraph(inner, item, size=7.5, color="1F2937", after=0.8)


def create_doc():
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(0.9)
    section.bottom_margin = Cm(0.75)
    section.left_margin = Cm(0.85)
    section.right_margin = Cm(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(1)
    run = title.add_run("北师版四年级下册  运算定律一页通")
    format_run(run, size=16, bold=True, color="0F3761")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    run = subtitle.add_run("交换律 · 结合律 · 分配律 | 规则、公式、易混口诀、专项练习")
    format_run(run, size=8.5, color="4B5563")

    main = doc.add_table(rows=1, cols=2)
    main.autofit = False
    set_table_borders(main, color="FFFFFF", size="0")
    left, right = main.rows[0].cells
    for cell, width in ((left, 9.5), (right, 8.6)):
        set_fixed_cell_width(cell, width)
        set_cell_margins(cell, top=0, start=75, bottom=0, end=75)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    add_small_box(
        left,
        "一、三大运算定律规则",
        ["先判断式子里有什么运算，再决定用哪条定律。", "只含加法或只含乘法：多想“换位置、加括号”。"],
        fill="EEF6FF",
        accent="1F4D78",
    )
    paragraph(left, "（一）加法运算定律", size=9, bold=True, color="0F3761", before=3, after=2)
    add_rule_card(left, "1. 加法交换律", "两个数相加，交换加数位置，和不变。", "a + b = b + a", "只换位置，运算符号、数字、运算级别都不变。")
    add_rule_card(left, "2. 加法结合律", "三个数相加，先加前两个，或先加后两个，和不变。", "(a + b) + c = a + (b + c)", "不换数字位置，只改变计算先后顺序，用括号分组。")
    paragraph(left, "（二）乘法运算定律", size=9, bold=True, color="0F3761", before=3, after=2)
    add_rule_card(left, "1. 乘法交换律", "两个数相乘，交换乘数位置，积不变。", "a × b = b × a", "同加法交换律，仅调换数字位置。")
    add_rule_card(left, "2. 乘法结合律", "三个数相乘，先乘前两个，或先乘后两个，积不变。", "(a × b) × c = a × (b × c)", "不换位置，只靠括号改变计算顺序，全是乘法。")
    add_rule_card(left, "3. 乘法分配律（重难点）", "两个数的和（或差）乘一个数，可分别相乘再相加（或相减）。", "(a + b) × c = a×c + b×c；a×c + b×c = (a + b)×c；(a - b) × c = a×c - b×c", "有乘有加/减，是唯一连接两种运算的定律，必考易错。", fill="FFF7ED")

    add_small_box(
        right,
        "二、易混区分口诀",
        ["1. 只有加法/只有乘法：用交换律、结合律。", "2. 既有乘法又有加/减：用乘法分配律。", "3. 看到相同因数：常常可以“提出来”。"],
        fill="FEF3C7",
        accent="92400E",
    )
    paragraph(right, "三、专项练习题", size=9, bold=True, color="0F3761", before=4, after=2)
    exercises = [
        ("基础：直接说出运用的定律", [
            "1. 37 + 46 = 46 + 37",
            "2. (25 + 18) + 82 = 25 + (18 + 82)",
            "3. 8 × 125 = 125 × 8",
            "4. (4 × 25) × 9 = 4 × (25 × 9)",
        ]),
        ("进阶：用简便方法计算", [
            "5. 125 × 32 × 25",
            "6. 48 × 99 + 48",
            "7. 101 × 36",
            "8. 25 × (40 + 4)",
        ]),
        ("易错拔高：判断并计算", [
            "9. 76 × 98 = 76 × (100 - 2)",
            "10. 35 × 27 + 65 × 27",
            "11. 99 × 45",
            "12. 88 + 125 + 12 + 75",
        ]),
    ]
    for heading, qs in exercises:
        paragraph(right, heading, size=8.2, bold=True, color="1F4D78", before=2.5, after=1)
        for q in qs:
            paragraph(right, q, size=7.8, color="111827", after=1.2)

    add_small_box(
        right,
        "答案速查",
        [
            "1 加法交换律；2 加法结合律；3 乘法交换律；4 乘法结合律。",
            "5 =125×8×4×25=1000×100=100000。",
            "6 =48×(99+1)=4800；7 =(100+1)×36=3636。",
            "8 =25×40+25×4=1100；9 =7600-152=7448。",
            "10 =(35+65)×27=2700；11=(100-1)×45=4455。",
            "12 =(88+12)+(125+75)=300。",
        ],
        fill="ECFDF5",
        accent="047857",
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(3)
    footer.paragraph_format.space_after = Pt(0)
    r = footer.add_run("打印建议：A4 纵向，缩放 100%，单页打印。")
    format_run(r, size=7, color="6B7280")

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    create_doc()
    print(DOCX_PATH.resolve())
